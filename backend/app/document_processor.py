import logging
from typing import Dict, Any, Tuple, Optional
import tempfile
import os
import io

# Import document processing libraries
import fitz  # PyMuPDF
import docx
import markdown

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    def __init__(self):
        """Initialize the document processor."""
        logger.info("Initializing document processor")
    
    def extract_text(self, content_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a document file.
        
        Args:
            content_bytes: The file content as bytes
            filename: Original filename with extension
            
        Returns:
            Tuple of (extracted text, metadata dictionary)
        """
        file_extension = self._get_file_extension(filename).lower()
        
        # Process based on file extension
        if file_extension == ".pdf":
            return self._extract_from_pdf(content_bytes)
        elif file_extension == ".docx":
            return self._extract_from_docx(content_bytes)
        elif file_extension == ".md":
            return self._extract_from_markdown(content_bytes)
        elif file_extension == ".txt":
            return self._extract_from_txt(content_bytes)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            raise ValueError(f"Unsupported file type: {file_extension}. Supported formats are PDF, DOCX, MD, and TXT.")
    
    def _get_file_extension(self, filename: str) -> str:
        """Get the file extension from a filename."""
        _, extension = os.path.splitext(filename)
        return extension
    
    def _extract_from_pdf(self, content_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a PDF file."""
        try:
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(content_bytes)
                temp_path = temp_file.name
            
            # Open the PDF with PyMuPDF
            doc = fitz.open(temp_path)
            
            # Extract text from all pages
            text = ""
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            # Extract metadata
            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
                "char_count": len(text)
            }
            
            # Clean up
            doc.close()
            os.unlink(temp_path)
            
            logger.info(f"Successfully extracted text from PDF: {len(text)} characters, {metadata['page_count']} pages")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def _extract_from_docx(self, content_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a DOCX file."""
        try:
            # Create a file-like object from bytes
            docx_file = io.BytesIO(content_bytes)
            
            # Open with python-docx
            doc = docx.Document(docx_file)
            
            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract metadata
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "char_count": len(text)
            }
            
            logger.info(f"Successfully extracted text from DOCX: {len(text)} characters")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def _extract_from_markdown(self, content_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a Markdown file."""
        try:
            # Decode bytes to string
            md_text = content_bytes.decode('utf-8')
            
            # Convert markdown to HTML and then extract text
            html = markdown.markdown(md_text)
            
            # Simple removal of HTML tags (a more sophisticated approach would use BeautifulSoup)
            # This is a simplified approach
            text = html.replace('<p>', '\n').replace('</p>', '\n')
            for tag in ['<h1>', '</h1>', '<h2>', '</h2>', '<h3>', '</h3>', '<strong>', '</strong>', '<em>', '</em>']:
                text = text.replace(tag, '')
            
            metadata = {
                "char_count": len(md_text),
                "line_count": md_text.count('\n') + 1
            }
            
            logger.info(f"Successfully extracted text from Markdown: {len(text)} characters")
            return md_text, metadata  # Return original markdown for better preservation
            
        except Exception as e:
            logger.error(f"Error extracting text from Markdown: {str(e)}")
            raise
    
    def _extract_from_txt(self, content_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a plain text file."""
        try:
            # Decode bytes to string
            text = content_bytes.decode('utf-8')
            
            metadata = {
                "char_count": len(text),
                "line_count": text.count('\n') + 1
            }
            
            logger.info(f"Successfully extracted text from TXT: {len(text)} characters")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise